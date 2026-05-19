"""Production document ingestion for Aptiro.

Turns an uploaded PDF / DOCX / TXT / Markdown file into clean,
section-friendly plain text plus a parse-metadata blob that records
format, page count, and a char-offset -> page map so the existing
parse_document / extract_claims pipeline can attach page numbers and
section/snippet provenance WITHOUT any change to how claims are derived.

This module deliberately does NOT change the claim model, provenance
rules, confidence scoring, or the human-approval gate. It only improves
how raw text is produced from real files.
"""
import io
import os
import re

MAX_UPLOAD_BYTES = int(os.getenv("APTIRO_MAX_UPLOAD_MB", "10")) * 1024 * 1024
SUPPORTED = {"pdf", "docx", "txt", "md", "markdown", "text"}


class IngestionError(Exception):
    pass


class UnsupportedFormat(IngestionError):
    pass


class ExtractionError(IngestionError):
    pass


class ExtractResult:
    def __init__(self, text, meta):
        self.text = text
        self.meta = meta


_BULLET_GLYPHS = "\u2022\u00b7\u25aa\u25e6\u2023\u2043\u204c\u2219"
_WS = re.compile(r"[ \t\u00a0]+")


def _ext(filename):
    return (os.path.splitext(filename or "")[1] or "").lower().lstrip(".")


def _normalize(text):
    """Collapse PDF/DOCX whitespace noise while preserving line structure
    and bullet markers (parse_document keys off newlines + bullet glyphs).
    """
    out = []
    for raw in (text or "").replace("\r\n", "\n").replace("\r", "\n").split(
            "\n"):
        line = _WS.sub(" ", raw).strip()
        # Drop pure table-rule lines from pandoc/markdown grid tables
        # (e.g. "+----+---+", "|====|", "---|---"): pure separator
        # punctuation, never resume content. Real piped data rows
        # (which contain letters/digits) are preserved.
        if line and not re.search(r"[0-9A-Za-z]", line) and \
                set(line) <= set("+-=|: "):
            continue
        # Normalize exotic bullet glyphs to a hyphen so the existing
        # _BULLET_PREFIX regex in app.parse_document picks them up.
        if line and line[0] in _BULLET_GLYPHS:
            line = "- " + line[1:].lstrip()
        out.append(line)
    # Drop runs of >2 blank lines.
    cleaned, blanks = [], 0
    for ln in out:
        if ln == "":
            blanks += 1
            if blanks > 1:
                continue
        else:
            blanks = 0
        cleaned.append(ln)
    return "\n".join(cleaned).strip()


def _txt(data):
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _markdown(data):
    """Markdown -> readable text: strip emphasis/heading markers but keep
    list items as '- ...' so they parse as bullets, and keep ATX headings
    on their own line so section detection still works.
    """
    raw = _txt(data)
    lines = []
    for ln in raw.split("\n"):
        s = ln.rstrip()
        s = re.sub(r"^\s{0,3}#{1,6}\s+", "", s)          # headings
        s = re.sub(r"^\s*[-*+]\s+", "- ", s)             # list items
        # Pandoc/DOCX-export artifacts: spans with attributes, nested
        # link/underline forms, and backslash-escaped punctuation. Real
        # resumes exported via pandoc are full of these; strip them so
        # claims read cleanly without changing any factual content.
        for _ in range(3):  # resolve nested [[x]{.underline}](url) forms
            new = re.sub(r"\[([^\]]*?)\]\{[^}]*\}", r"\1", s)   # [x]{.attr}
            new = re.sub(r"\[([^\]]+?)\]\((?:[^)]*)\)", r"\1", new)  # links
            if new == s:
                break
            s = new
        s = re.sub(r"\{[.#][^}]*\}", "", s)              # leftover {.attr}
        s = re.sub(r"\\([\\`*_{}\[\]()#+\-.!~|>$@&])", r"\1", s)  # md esc
        s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)           # bold
        s = re.sub(r"(?<!\*)\*(?!\*)(.+?)\*", r"\1", s)  # italic
        s = re.sub(r"`([^`]+)`", r"\1", s)               # inline code
        s = re.sub(r"\[(.+?)\]\((.+?)\)", r"\1", s)      # any leftover links
        s = re.sub(r"\\+$", "", s)                       # pandoc hard break
        s = re.sub(r"^\s*[*_]\s*$", "", s)               # lone emphasis line
        s = re.sub(r"^[*_]+(?=\S)|(?<=\S)[*_]+$", "", s)  # dangling emphasis
        lines.append(s)
    return "\n".join(lines)


def _docx(data):
    try:
        from docx import Document
    except Exception as e:                                # pragma: no cover
        raise ExtractionError(
            "python-docx is required for .docx ingestion (%s)." % e)
    try:
        doc = Document(io.BytesIO(data))
    except Exception as e:
        raise ExtractionError("Could not read the .docx file: %s" % e)
    parts = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if not t:
            parts.append("")
            continue
        style = (p.style.name or "").lower() if p.style else ""
        if "list" in style or "bullet" in style:
            t = "- " + t
        parts.append(t)
    # Tables (skills grids etc.) -> one line per row.
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _pdf(data):
    """Text + a char-offset -> page map (1-indexed) so each claim can
    record the page it came from. Uses pdfplumber (text-faithful)."""
    try:
        import pdfplumber
    except Exception as e:                                # pragma: no cover
        raise ExtractionError(
            "pdfplumber is required for .pdf ingestion (%s)." % e)
    try:
        pdf = pdfplumber.open(io.BytesIO(data))
    except Exception as e:
        raise ExtractionError("Could not open the PDF: %s" % e)
    chunks = []
    page_map = []
    offset = 0
    n_pages = 0
    with pdf:
        n_pages = len(pdf.pages)
        for i, page in enumerate(pdf.pages, start=1):
            txt = page.extract_text() or ""
            page_map.append([offset, i])
            chunks.append(txt)
            offset += len(txt) + 1
    return "\n".join(chunks), page_map, n_pages


def extract(filename, data):
    """Dispatch on extension. Returns ExtractResult(text, meta).

    Raises UnsupportedFormat / ExtractionError; the caller maps these to
    415 / 422 so the user gets a clear, non-crashing message.
    """
    if not data:
        raise ExtractionError("Empty file.")
    ext = _ext(filename)
    meta = {"filename": filename, "format": ext or "unknown",
            "bytes": len(data)}

    if ext == "pdf":
        text, page_map, n_pages = _pdf(data)
        text = _normalize(text)
        meta.update({"format": "pdf", "pages": n_pages,
                     "page_map": page_map})
    elif ext == "docx":
        text = _normalize(_docx(data))
        meta["format"] = "docx"
    elif ext in ("md", "markdown"):
        text = _normalize(_markdown(data))
        meta["format"] = "markdown"
    elif ext in ("txt", "text", ""):
        text = _normalize(_txt(data))
        meta["format"] = "txt"
    elif ext == "doc":
        raise UnsupportedFormat(
            "Legacy .doc is not supported - re-save as .docx or PDF.")
    else:
        raise UnsupportedFormat(
            "Unsupported file type '.%s'. Supported: PDF, DOCX, TXT, "
            "Markdown." % ext)

    meta["chars"] = len(text)
    meta["lines"] = text.count("\n") + 1 if text else 0
    return ExtractResult(text, meta)
